import docx


d = docx.Document('.\\Documents\\demo.docx')

# print(d.paragraphs[0].text)

p = d.paragraphs[1]

# Runs are start from the beginning of the line then stops -
# when the style of the paragraph/sentence changes.

# Ex:  Default Style text -> Bold -> Default -> Italic
# Each time the style changes it creates a new run
p.runs[0].text # A plain paragraph having some 
p.runs[1].text # bold
p.runs[2].text # and some
p.runs[3].text # italic

# You can also check if the run is in italic or bold by using
# When run they will return a true or false value
# p.runs[1].bold 
# p.runs[3].italic 

# You can also set runs property to true or false
p.runs[3].underline = True
p.runs[3].text = 'italic and underlined'

# Saving/Creating new document
d.save('.\\Documents\\demo2.docx')

# You can check the style of the paragraph by - 
# p.style
# You can also changed the style by setting it = 
# to the style that you like
# p.style = 'Title'

p.style = 'Title'
d.save('.\\Documents\\demo3.docx')


d = docx.Document()

d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph('This is another paragraph.')

d.save('.\\Documents\\demo4.docx')

# Adding paragraphs to existing words
p = d.paragraphs[0]

p.add_run('This is a new run.')
p.runs[1].bold = True

d.save('.\\Documents\\demo5.docx')
